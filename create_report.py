from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Add title
title = doc.add_heading('Project Report: Digital Wellbeing Java Desktop Application', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add sections and content
def add_section(doc, title, content):
    doc.add_heading(title, level=1)
    for paragraph_text in content:
        p = doc.add_paragraph(paragraph_text)
        p.style.font.size = Pt(11)

# Section 1: Introduction
introduction = [
    "In today's fast-paced digital world, maintaining a healthy balance between technology use and personal well-being has become increasingly important. The Digital Wellbeing Java Desktop Application was developed to address this need by providing users with tools to monitor and manage their digital habits. This report outlines the objectives, design, implementation, and outcomes of the project."
]
add_section(doc, "1. Introduction", introduction)

# Section 2: Project Objectives
objectives = [
    "The primary objectives of the Digital Wellbeing application are:",
    "• Monitor Usage: Track the time spent on various digital activities.",
    "• Promote Healthy Habits: Encourage users to take breaks and reduce screen time.",
    "• User-Friendly Interface: Provide an intuitive and accessible desktop application.",
    "• Cross-Platform Compatibility: Ensure the application runs smoothly on different operating systems."
]
add_section(doc, "2. Project Objectives", objectives)

# Section 3: System Design
system_design = [
    "3.1 Architecture",
    "The application follows a modular architecture, separating the user interface, business logic, and data management layers. This design enhances maintainability and scalability.",
    "• User Interface (UI): Built using Java Swing/JavaFX, providing a responsive and interactive experience.",
    "• Business Logic: Handles core functionalities such as tracking usage, setting reminders, and generating reports.",
    "• Data Management: Stores user data locally, ensuring privacy and quick access.",
    "",
    "3.2 Key Features",
    "• Activity Tracking: Monitors the duration of application usage.",
    "• Break Reminders: Notifies users to take regular breaks.",
    "• Usage Reports: Generates daily/weekly reports to help users analyze their digital habits.",
    "• Customizable Settings: Allows users to set their own goals and reminder intervals."
]
add_section(doc, "3. System Design", system_design)

# Section 4: Implementation
implementation = [
    "4.1 Development Environment",
    "• Programming Language: Java",
    "• Build Tool: Maven",
    "• IDE: IntelliJ IDEA",
    "• JAR Packaging: The application is packaged as a runnable JAR file for easy distribution.",
    "",
    "4.2 Running the Application",
    "A batch script (run_app.bat) is provided to simplify launching the application. The script executes the following commands:",
    "@echo off",
    "java -jar target\\DigitalWellbeing-1.0-SNAPSHOT.jar",
    "pause",
    "This ensures that users can start the application with a double-click, and view any output or errors before the window closes.",
    "",
    "4.3 Challenges Faced",
    "• Cross-Platform Issues: Ensuring consistent behavior across different operating systems required careful handling of file paths and system resources.",
    "• User Experience: Designing an interface that is both functional and aesthetically pleasing was a key focus.",
    "• Data Privacy: Storing user data securely and locally to maintain privacy."
]
add_section(doc, "4. Implementation", implementation)

# Section 5: Testing and Evaluation
testing = [
    "The application was tested on multiple systems to ensure reliability and performance. User feedback was collected to refine the interface and improve usability. Key evaluation criteria included:",
    "• Stability: The application runs without crashes or major bugs.",
    "• Performance: Resource usage is minimal, ensuring smooth operation.",
    "• User Satisfaction: Users found the application helpful in managing their digital habits."
]
add_section(doc, "5. Testing and Evaluation", testing)

# Section 6: Conclusion
conclusion = [
    "The Digital Wellbeing Java Desktop Application successfully meets its objectives by providing users with practical tools to monitor and improve their digital habits. The project demonstrates the effective use of Java for desktop application development and highlights the importance of user-centric design."
]
add_section(doc, "6. Conclusion", conclusion)

# Section 7: Future Work
future_work = [
    "Potential enhancements for future versions include:",
    "• Cloud Sync: Allowing users to back up and sync their data across devices.",
    "• Advanced Analytics: Providing deeper insights into usage patterns.",
    "• Integration with Other Platforms: Extending support to mobile devices."
]
add_section(doc, "7. Future Work", future_work)

# Section 8: References
references = [
    "• Java Documentation: https://docs.oracle.com/javase/",
    "• JavaFX/Swing Documentation",
    "• Maven Documentation"
]
add_section(doc, "8. References", references)

# Save the document
doc.save('DigitalWellbeing_Project_Report.docx')

print("Project report has been created as 'DigitalWellbeing_Project_Report.docx'") 